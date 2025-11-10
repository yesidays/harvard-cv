"""CV generation services for PDF and DOCX formats."""
import os
from typing import Dict, Any
from io import BytesIO
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..core.config import settings


class CVGenerator:
    """CV generator for Harvard format."""

    def __init__(self):
        """Initialize CV generator with template environment."""
        self.template_env = Environment(
            loader=FileSystemLoader(settings.TEMPLATES_DIR)
        )

    def _prepare_cv_data(self, cv_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare CV data for template rendering."""
        # Sort education and experience by date (most recent first)
        education = cv_data.get("education", [])
        experience = cv_data.get("experience", [])

        # Sort by start_date in reverse
        education_sorted = sorted(
            education,
            key=lambda x: x.get("start_date", ""),
            reverse=True,
        )
        experience_sorted = sorted(
            experience,
            key=lambda x: x.get("start_date", ""),
            reverse=True,
        )

        return {
            "profile": cv_data.get("profile", {}),
            "education": education_sorted,
            "experience": experience_sorted,
            "certifications": cv_data.get("certifications", []),
            "projects": cv_data.get("projects", []),
            "skills": cv_data.get("skills"),
        }

    def generate_html(self, cv_data: Dict[str, Any]) -> str:
        """Generate HTML preview of CV."""
        template = self.template_env.get_template("harvard_cv.html")
        prepared_data = self._prepare_cv_data(cv_data)
        return template.render(**prepared_data)

    def generate_pdf(self, cv_data: Dict[str, Any]) -> bytes:
        """Generate PDF from CV data."""
        html_content = self.generate_html(cv_data)
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes

    def generate_docx(self, cv_data: Dict[str, Any]) -> BytesIO:
        """Generate DOCX from CV data."""
        prepared_data = self._prepare_cv_data(cv_data)
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        profile = prepared_data["profile"]

        # Header - Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(
            f"{profile.get('first_name', '')} {profile.get('last_name', '')}"
        )
        name_run.font.name = "Georgia"
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = [profile.get("email", "")]
        if profile.get("phone"):
            contact_parts.append(profile["phone"])
        if profile.get("linkedin"):
            contact_parts.append(profile["linkedin"])
        if profile.get("location"):
            contact_parts.append(profile["location"])

        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.runs[0]
        contact_run.font.size = Pt(10)

        # Add horizontal line
        doc.add_paragraph("_" * 80)

        # Summary
        if profile.get("summary"):
            doc.add_paragraph(profile["summary"])

        # Education
        if prepared_data["education"]:
            self._add_section_title(doc, "EDUCATION")
            for edu in prepared_data["education"]:
                self._add_education_entry(doc, edu)

        # Experience
        if prepared_data["experience"]:
            self._add_section_title(doc, "EXPERIENCE")
            for exp in prepared_data["experience"]:
                self._add_experience_entry(doc, exp)

        # Projects
        if prepared_data["projects"]:
            self._add_section_title(doc, "PROJECTS")
            for project in prepared_data["projects"]:
                self._add_project_entry(doc, project)

        # Certifications
        if prepared_data["certifications"]:
            self._add_section_title(doc, "CERTIFICATIONS")
            for cert in prepared_data["certifications"]:
                self._add_certification_entry(doc, cert)

        # Skills
        if prepared_data["skills"]:
            self._add_section_title(doc, "SKILLS")
            self._add_skills(doc, prepared_data["skills"])

        # Save to BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes

    def _add_section_title(self, doc, title: str):
        """Add section title to document."""
        para = doc.add_paragraph(title)
        run = para.runs[0]
        run.font.name = "Georgia"
        run.font.size = Pt(13)
        run.font.bold = True
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _add_education_entry(self, doc, edu: Dict[str, Any]):
        """Add education entry to document."""
        # Institution and dates
        para = doc.add_paragraph()
        institution_run = para.add_run(edu.get("institution", ""))
        institution_run.font.bold = True

        # Add dates on the same line (right-aligned would need table)
        dates = self._format_dates(edu.get("start_date"), edu.get("end_date"))
        if dates:
            para.add_run(f"\t{dates}")

        # Degree
        degree_text = edu.get("degree", "")
        if edu.get("location"):
            degree_text += f", {edu['location']}"
        degree_para = doc.add_paragraph(degree_text)
        degree_para.runs[0].font.italic = True

        # Details
        if edu.get("details"):
            for detail in edu["details"]:
                doc.add_paragraph(detail, style="List Bullet")

        doc.add_paragraph()  # Space after entry

    def _add_experience_entry(self, doc, exp: Dict[str, Any]):
        """Add experience entry to document."""
        # Company and role
        para = doc.add_paragraph()
        title_run = para.add_run(
            f"{exp.get('company', '')} — {exp.get('role', '')}"
        )
        title_run.font.bold = True

        # Dates
        dates = self._format_dates(exp.get("start_date"), exp.get("end_date"))
        if dates:
            para.add_run(f"\t{dates}")

        # Location
        if exp.get("location"):
            loc_para = doc.add_paragraph(exp["location"])
            loc_para.runs[0].font.italic = True

        # Bullets
        if exp.get("bullets"):
            for bullet in exp["bullets"]:
                doc.add_paragraph(bullet, style="List Bullet")

        doc.add_paragraph()  # Space after entry

    def _add_certification_entry(self, doc, cert: Dict[str, Any]):
        """Add certification entry to document."""
        para = doc.add_paragraph()
        name_run = para.add_run(cert.get("name", ""))
        name_run.font.bold = True

        # Date
        if cert.get("date"):
            para.add_run(f"\t{cert['date']}")

        # Issuer and credential
        issuer_text = cert.get("issuer", "")
        if cert.get("credential_id"):
            issuer_text += f", Credential: {cert['credential_id']}"

        if issuer_text:
            issuer_para = doc.add_paragraph(issuer_text)
            issuer_para.runs[0].font.italic = True

    def _add_project_entry(self, doc, project: Dict[str, Any]):
        """Add project entry to document."""
        # Project name
        para = doc.add_paragraph()
        name_run = para.add_run(project.get("name", ""))
        name_run.font.bold = True

        if project.get("url"):
            para.add_run(f" ({project['url']})")

        # Impact
        if project.get("impact"):
            doc.add_paragraph(project["impact"])

        # Technologies
        if project.get("technologies"):
            tech_text = "Technologies: " + ", ".join(project["technologies"])
            tech_para = doc.add_paragraph(tech_text)
            tech_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # Space after entry

    def _add_skills(self, doc, skills: Dict[str, Any]):
        """Add skills section to document."""
        if skills.get("languages"):
            para = doc.add_paragraph()
            label_run = para.add_run("Languages: ")
            label_run.font.bold = True
            para.add_run(", ".join(skills["languages"]))

        if skills.get("tools"):
            para = doc.add_paragraph()
            label_run = para.add_run("Tools: ")
            label_run.font.bold = True
            para.add_run(", ".join(skills["tools"]))

        if skills.get("methods"):
            para = doc.add_paragraph()
            label_run = para.add_run("Methodologies: ")
            label_run.font.bold = True
            para.add_run(", ".join(skills["methods"]))

    def _format_dates(self, start_date: str, end_date: str = None) -> str:
        """Format date range."""
        if not start_date:
            return ""
        if end_date:
            return f"{start_date} – {end_date}"
        return f"{start_date} – Present"

    def get_filename(self, profile: Dict[str, Any], format: str) -> str:
        """Generate filename for CV."""
        last_name = profile.get("last_name", "CV")
        first_name = profile.get("first_name", "")
        return f"{last_name}_{first_name}_CV_Harvard.{format}"


# Singleton instance
cv_generator = CVGenerator()
